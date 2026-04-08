from flask import Flask, render_template, request, jsonify, send_file
import json, os, re, time, threading, zipfile, io
from pathlib import Path
from datetime import datetime
import requests
from bs4 import BeautifulSoup
import anthropic
from docx import Document
from docx.shared import Pt
import psycopg2
from psycopg2.extras import RealDictCursor

app = Flask(__name__)
# Gebruik /data als persistent volume (Railway), anders lokale output map
if Path("/data").exists():
    OUTPUT_DIR = Path("/data")
else:
    OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)
STATE_FILE = OUTPUT_DIR / ".state.json"
SEEN_FILE = OUTPUT_DIR / ".seen.json"

SYSTEM_PROMPT = """Je bent een senior Microsoft 365 / Modern Workplace engineer die Message Center items analyseert.
Schrijf ALTIJD in het Nederlands. Geen em-dash. Geen "ten eerste/tweede". Omschrijving zonder risico/impact.
Geef ALLEEN pure JSON terug - geen markdown, geen backticks.

{"mcId":"MC1234567","title":"[Platform] Titel [MC1234567]","platform":"platform","roadmapId":"id of null","roadmapUrl":"https://www.microsoft.com/microsoft-365/roadmap","plannerTask":"[Platform] Titel [MC1234567]","planning":["Targeted Release: ...","Algemeen beschikbaar: ..."],"oneLiner":"Max 2 zinnen geschikt als opmerking in Planner. Zakelijk en concreet.","omschrijvingIntro":"tekst","omschrijvingBullets":["punt1","punt2"],"omschrijvingSlot":"tekst of lege string","impactOrganisaties":"laag/gemiddeld/hoog - toelichting","impactTechnisch":"tekst","impactFunctioneel":"tekst","impactBeheer":["actie1","actie2"],"relevantieSCore":3,"relevantieUitleg":"Max 1 zin waarom dit item relevant of minder relevant is.","links":[{"label":"Microsoft Learn - naam","url":"https://..."},{"label":"Microsoft Message Center - MC1234567","url":null}],"geenSpecifiekeLearnPagina":false}

De relevantieScore is een getal van 1 tot 5:
1 = Nauwelijks relevant (bijv. Dynamics 365 specifiek)
2 = Beperkt relevant
3 = Gemiddeld relevant
4 = Relevant voor de meeste M365 organisaties
5 = Zeer relevant, actie vereist"""

progress = {"total": 0, "done": 0, "current": "", "running": False, "errors": [], "new_analyzed": []}

def load_state():
    try:
        conn = get_db()
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute('SELECT * FROM analyses')
        rows = cur.fetchall()
        cur.close(); conn.close()
        return {row['mc_id']: {
            'title': row['title'],
            'filename': row['filename'],
            'analyzed_at': row['analyzed_at'],
            'analysis': row['analysis']
        } for row in rows}
    except Exception as e:
        print(f"load_state fout: {e}")
        return {}

def save_analysis(mc_id, title, filename, analyzed_at, analysis):
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute('''INSERT INTO analyses (mc_id, title, filename, analyzed_at, analysis)
            VALUES (%s, %s, %s, %s, %s)
            ON CONFLICT (mc_id) DO UPDATE SET
            title=EXCLUDED.title, filename=EXCLUDED.filename,
            analyzed_at=EXCLUDED.analyzed_at, analysis=EXCLUDED.analysis''',
            (mc_id, title, filename, analyzed_at, json.dumps(analysis)))
        conn.commit()
        cur.close(); conn.close()
    except Exception as e:
        print(f"save_analysis fout: {e}")

def load_seen():
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute('SELECT mc_id FROM seen_items')
        rows = cur.fetchall()
        cur.close(); conn.close()
        return set(row[0] for row in rows)
    except:
        return set()

def save_seen(seen):
    try:
        conn = get_db()
        cur = conn.cursor()
        for mc_id in seen:
            cur.execute('INSERT INTO seen_items (mc_id) VALUES (%s) ON CONFLICT DO NOTHING', (mc_id,))
        conn.commit()
        cur.close(); conn.close()
    except Exception as e:
        print(f"save_seen fout: {e}")

def fetch_mc_list(count):
    resp = requests.get("https://mc.merill.net", timeout=15)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    items = []
    for row in soup.select("table tr"):
        cells = row.select("td")
        if len(cells) < 4: continue
        mc_id = cells[0].get_text(strip=True)
        if not mc_id.startswith("MC"): continue
        items.append({"id": mc_id, "title": cells[1].get_text(strip=True),
                      "service": cells[2].get_text(strip=True),
                      "lastUpdated": cells[3].get_text(strip=True),
                      "url": f"https://mc.merill.net/message/{mc_id}"})
        if len(items) >= count: break
    return items

def fetch_item_text(item):
    resp = requests.get(item["url"], timeout=20)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    main = soup.find("main") or soup.find("article") or soup.body
    text = main.get_text(separator="\n", strip=True)
    return f"Message ID: {item['id']}\nTitle: {item['title']}\nService: {item['service']}\n\n{text[:8000]}"

def fetch_item_images(url):
    try:
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        main = soup.find("main") or soup.find("article") or soup.body
        imgs = []
        for i, img in enumerate(main.find_all("img"), 1):
            src = img.get("src") or img.get("data-src")
            if src and not src.startswith("data:") and len(src) > 10:
                if src.startswith("/"): src = "https://mc.merill.net" + src
                alt = img.get("alt") or f"Afbeelding {i}"
                imgs.append({"url": src, "alt": alt, "index": i})
        return imgs
    except:
        return []

def analyze(client, text):
    msg = client.messages.create(
        model="claude-sonnet-4-6", max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": text}],
        timeout=60.0)
    raw = msg.content[0].text.strip()
    if raw.startswith("```"):
        raw = re.sub(r'^```(?:json)?\n?', '', raw)
        raw = re.sub(r'\n?```$', '', raw)
    return json.loads(raw)

def build_docx(a, path):
    doc = Document()
    def bp(t): p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(11)
    def np(t):
        if not t: doc.add_paragraph(); return
        p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(11)
    def lp(l, v):
        p = doc.add_paragraph(); r1 = p.add_run(l); r1.bold = True; r1.font.size = Pt(11)
        r2 = p.add_run(v or ""); r2.font.size = Pt(11)
    def bl(t):
        p = doc.add_paragraph(style="List Bullet"); r = p.add_run(t); r.font.size = Pt(11)

    p = doc.add_paragraph(); r = p.add_run(a.get("title","")); r.bold = True; r.font.size = Pt(12)
    doc.add_paragraph()
    bp("Platform:"); np(a.get("platform",""))
    doc.add_paragraph()
    bp("Link naar Microsoft (Roadmap ID + URL):")
    np(f"Roadmap ID: {a.get('roadmapId') or 'niet van toepassing'}")
    if a.get("roadmapUrl"): np(a["roadmapUrl"])
    doc.add_paragraph()
    bp("Link naar Teams taak:"); np(f"Planner - {a.get('plannerTask','')}")
    doc.add_paragraph()
    bp("Planning:")
    for l in (a.get("planning") or []): np(l)
    doc.add_paragraph()
    bp("Omschrijving wijziging:")
    if a.get("omschrijvingIntro"): np(a["omschrijvingIntro"])
    if a.get("omschrijvingBullets"):
        doc.add_paragraph()
        for b in a["omschrijvingBullets"]: bl(b)
    if a.get("omschrijvingSlot"): doc.add_paragraph(); np(a["omschrijvingSlot"])
    doc.add_paragraph()
    bp("Impactanalyse:")
    lp("Impact voor organisaties: ", a.get("impactOrganisaties",""))
    lp("Technische impact: ", a.get("impactTechnisch",""))
    lp("Functionele impact: ", a.get("impactFunctioneel",""))
    bp("Wijzigingen in beheer of gedrag:")
    for b in (a.get("impactBeheer") or []): bl(b)
    doc.add_paragraph()
    bp("Links:")
    if a.get("geenSpecifiekeLearnPagina"):
        np("Geen specifieke Microsoft Learn-pagina voor deze update gevonden. Hieronder de meest relevante officiële bronnen.")
        doc.add_paragraph()
    for link in (a.get("links") or []):
        bp(f"{link.get('label','')}:")
        np(link.get("url") or "Microsoft Message Center")
        doc.add_paragraph()
    doc.save(str(path))

def send_teams_notification(webhook_url, new_items):
    if not webhook_url or not new_items: return
    items_text = "\n".join([f"- **{i['mcId']}** {i['title']} (score: {i.get('relevantieSCore','?')}/5)" for i in new_items[:10]])
    payload = {
        "@type": "MessageCard", "@context": "http://schema.org/extensions",
        "themeColor": "0078D4", "summary": f"{len(new_items)} nieuwe MC analyses gereed",
        "sections": [{"activityTitle": f"MC Analyzer: {len(new_items)} nieuwe analyses",
                      "activitySubtitle": "Microsoft 365 Message Center",
                      "activityText": f"De volgende items zijn geanalyseerd:\n\n{items_text}", "markdown": True}]
    }
    try: requests.post(webhook_url, json=payload, timeout=10)
    except Exception as e: print(f"Teams notificatie mislukt: {e}")

def run_analysis(api_key, items, force, webhook_url=""):
    global progress
    client = anthropic.Anthropic(api_key=api_key)
    progress["running"] = True
    progress["total"] = len(items)
    progress["done"] = 0
    progress["errors"] = []
    progress["new_analyzed"] = []

    for item in items:
        if not progress["running"]: break  # Stop als reset is gedrukt
        mc_id = item["id"]
        progress["current"] = mc_id
        if not force and mc_id in state:
            progress["done"] += 1
            continue
        try:
            text = fetch_item_text(item)
            time.sleep(1)  # Vertraging om rate limiting te voorkomen
            result = analyze(client, text)
            time.sleep(2)  # Extra vertraging tussen items
            safe_title = re.sub(r'[\\/*?:"<>|]', '', result.get("title", mc_id))[:120]
            filename = f"{safe_title}.docx"
            docx_path = OUTPUT_DIR / filename
            build_docx(result, docx_path)
            save_analysis(mc_id, item["title"], filename, datetime.now().isoformat(), result)
            progress["new_analyzed"].append({
                "mcId": mc_id, "title": result.get("title", item["title"]),
                "relevantieSCore": result.get("relevantieSCore", 3)
            })
        except Exception as e:
            progress["errors"].append(f"{mc_id}: {str(e)}")
            print(f"Fout bij {mc_id}: {e}")
        progress["done"] += 1

    if webhook_url and progress["new_analyzed"]:
        send_teams_notification(webhook_url, progress["new_analyzed"])

    progress["running"] = False
    progress["current"] = ""

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/items")
def get_items():
    count = int(request.args.get("count", 50))
    try:
        items = fetch_mc_list(count)
        state = load_state()
        seen = load_seen()
        new_ids = []
        for item in items:
            item["status"] = "done" if item["id"] in state else "new"
            item["isNew"] = item["id"] not in seen
            if item["id"] not in seen: new_ids.append(item["id"])
            if item["id"] in state and state[item["id"]].get("analysis"):
                item["relevantieSCore"] = state[item["id"]]["analysis"].get("relevantieSCore", None)
                item["analyzedTitle"] = state[item["id"]]["analysis"].get("title", None)
        seen.update(i["id"] for i in items)
        save_seen(seen)
        return jsonify({"ok": True, "items": items, "newCount": len(new_ids)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})

@app.route("/api/analyze", methods=["POST"])
def start_analyze():
    global progress
    if progress["running"]:
        return jsonify({"ok": False, "error": "Al bezig"})
    data = request.json
    api_key = data.get("api_key") or os.environ.get("ANTHROPIC_API_KEY", "")
    items = data.get("items", [])
    force = data.get("force", False)
    webhook_url = data.get("webhook_url", "") or os.environ.get("TEAMS_WEBHOOK_URL", "")
    if not api_key: return jsonify({"ok": False, "error": "Geen API key"})
    t = threading.Thread(target=run_analysis, args=(api_key, items, force, webhook_url))
    t.daemon = True
    t.start()
    return jsonify({"ok": True})

@app.route("/api/reset", methods=["POST"])
def reset_progress():
    global progress
    progress = {"total": 0, "done": 0, "current": "", "running": False, "errors": [], "new_analyzed": []}
    return jsonify({"ok": True})

@app.route("/api/progress")
def get_progress():
    return jsonify(progress)

@app.route("/api/analyses")
def get_analyses():
    return jsonify({"ok": True, "analyses": load_state()})

@app.route("/api/download/<mc_id>")
def download_file(mc_id):
    state = load_state()
    entry = state.get(mc_id, {})
    if not entry:
        return "Niet gevonden", 404
    analysis = entry.get("analysis", {})
    filename = entry.get("filename", f"{mc_id}_analyse.docx")
    # Bouw docx opnieuw als het niet meer bestaat
    path = OUTPUT_DIR / filename
    if not path.exists():
        try:
            build_docx(analysis, path)
        except Exception as e:
            return f"Fout bij aanmaken document: {e}", 500
    return send_file(str(path), as_attachment=True, download_name=filename)

@app.route("/api/download-zip", methods=["POST"])
def download_zip():
    ids = request.json.get("ids", [])
    if not ids: return "Geen items opgegeven", 400
    state = load_state()
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for mc_id in ids:
            entry = state.get(mc_id, {})
            if not entry: continue
            filename = entry.get("filename", f"{mc_id}_analyse.docx")
            path = OUTPUT_DIR / filename
            if not path.exists():
                try: build_docx(entry.get("analysis", {}), path)
                except: continue
            if path.exists(): zf.write(path, filename)
    buf.seek(0)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    return send_file(buf, as_attachment=True,
                     download_name=f"MC_analyses_{timestamp}.zip",
                     mimetype="application/zip")

@app.route("/api/images/<mc_id>")
def get_images(mc_id):
    url = f"https://mc.merill.net/message/{mc_id}"
    images = fetch_item_images(url)
    return jsonify({"ok": True, "images": images})

@app.route("/api/settings", methods=["GET", "POST"])
def settings():
    if request.method == "POST": return jsonify({"ok": True})
    return jsonify({"api_key": os.environ.get("ANTHROPIC_API_KEY", ""),
                    "count": "50",
                    "webhook_url": os.environ.get("TEAMS_WEBHOOK_URL", "")})

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5001))
    print(f"\n MC Analyzer gestart op http://localhost:{port}\n")
    app.run(debug=False, host="0.0.0.0", port=port)
